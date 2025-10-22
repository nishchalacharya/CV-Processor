# import os
# from PyPDF2 import PdfReader
# from docx import Document

# def extract_text_from_pdf(file_path):
#     reader = PdfReader(file_path)
#     text = ""
#     for page in reader.pages:
#         text += page.extract_text() + "\n"
#     return text

# def extract_text_from_docx(file_path):
#     doc = Document(file_path)
#     text = "\n".join([para.text for para in doc.paragraphs])
#     return text

# def extract_text(file_path):
#     ext = os.path.splitext(file_path)[1].lower()
#     if ext == ".pdf":
#         return extract_text_from_pdf(file_path)
#     elif ext == ".docx":
#         return extract_text_from_docx(file_path)
#     elif ext == ".txt":
#         with open(file_path, 'r', encoding='utf-8') as f:
#             return f.read()
#     else:
#         raise ValueError(f"Unsupported file format: {ext}")












import fitz  # PyMuPDF
from docx import Document
import pytesseract
from pdf2image import convert_from_path
import re

class TextExtractor:
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt']
    
    def extract_text(self, file_path: str, file_type: str) -> str:
        """Extract text from document with universal format handling"""
        try:
            if file_type == 'pdf':
                return self._parse_pdf(file_path)
            elif file_type == 'docx':
                return self._parse_docx(file_path)
            elif file_type == 'txt':
                return self._parse_txt(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_type}")
        except Exception as e:
            print(f"Error parsing document: {e}")
            return ""
    
    def _parse_pdf(self, file_path: str) -> str:
        """Extract text from PDF with OCR fallback"""
        text = ""
        try:
            doc = fitz.open(file_path)
            for page in doc:
                text += page.get_text()
            
            if len(self._clean_text(text)) < 200:
                text = self._ocr_pdf(file_path)
            
            doc.close()
        except Exception as e:
            print(f"PDF parsing error: {e}")
            text = self._ocr_pdf(file_path)
        
        return self._clean_text(text)
    
    def _ocr_pdf(self, file_path: str) -> str:
        """Use OCR for scanned PDFs"""
        try:
            images = convert_from_path(file_path, dpi=300)
            text = ""
            for image in images:
                text += pytesseract.image_to_string(image)
            return self._clean_text(text)
        except Exception as e:
            print(f"OCR error: {e}")
            return ""
    
    def _parse_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return self._clean_text(text)
    
    def _parse_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            return self._clean_text(file.read())
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'[^\w\s.,!?;:()\-@]', '', text)
        return text.strip()
